import io
import os
import json
import uuid
import time
import subprocess
import tempfile
from datetime import datetime, timedelta
from typing import List, Dict, Any

import pandas as pd
import streamlit as st
from docx import Document

# -------------------------------
# FFmpeg: try portable binary via imageio-ffmpeg first
# -------------------------------
FFMPEG_PATH = None
try:
    import imageio_ffmpeg  # provided via requirements; ok to fail if system ffmpeg exists
    FFMPEG_PATH = imageio_ffmpeg.get_ffmpeg_exe()
except Exception:
    FFMPEG_PATH = None  # fall back to system ffmpeg on PATH

# -------------------------------
# Constants & paths
# -------------------------------
SUPPORTED_TYPES = ["mp4", "mp3", "wav", "m4a", "mov", "mkv", "avi"]
REPO_DIR = "repo"              # folder that stores all saved transcripts
INDEX_PATH = os.path.join(REPO_DIR, "index.json")  # registry file

# -------------------------------
# Small utilities
# -------------------------------
def ensure_repo():
    os.makedirs(REPO_DIR, exist_ok=True)
    if not os.path.exists(INDEX_PATH):
        with open(INDEX_PATH, "w", encoding="utf-8") as f:
            json.dump({"items": []}, f, ensure_ascii=False, indent=2)

def load_index() -> Dict[str, Any]:
    ensure_repo()
    with open(INDEX_PATH, "r", encoding="utf-8") as f:
        return json.load(f)

def save_index(data: Dict[str, Any]):
    with open(INDEX_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def new_id() -> str:
    return uuid.uuid4().hex[:12]

def human_dt(ts: float) -> str:
    return datetime.fromtimestamp(ts).strftime("%Y-%m-%d %H:%M")

def format_timestamp(seconds: float) -> str:
    td = timedelta(seconds=float(seconds))
    hours, remainder = divmod(td.seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    milliseconds = int(td.microseconds / 1000)
    hours += td.days * 24
    return f"{hours:02d}:{minutes:02d}:{seconds:02d},{milliseconds:03d}"

# -------------------------------
# Media helpers
# -------------------------------
def extract_audio_to_wav(in_bytes: bytes, out_path: str):
    """
    Convert any audio/video bytes to 16kHz mono WAV using ffmpeg.
    Uses portable imageio-ffmpeg binary if present; otherwise relies on system ffmpeg on PATH.
    """
    in_file = tempfile.NamedTemporaryFile(delete=False, suffix=".input")
    in_file.write(in_bytes)
    in_file.flush()
    in_file.close()

    ffmpeg_bin = FFMPEG_PATH or "ffmpeg"
    cmd = [ffmpeg_bin, "-y", "-i", in_file.name, "-ac", "1", "-ar", "16000", "-f", "wav", out_path]
    try:
        subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except FileNotFoundError:
        raise RuntimeError(
            "FFmpeg binary not found. Install system FFmpeg (e.g. winget install Gyan.FFmpeg) "
            "or keep 'imageio-ffmpeg' in requirements so a portable binary is available."
        )
    finally:
        try:
            os.unlink(in_file.name)
        except OSError:
            pass

# -------------------------------
# Transcription helpers
# -------------------------------
def segments_to_plaintext(segments) -> str:
    return "\n".join(seg.text.strip() for seg in segments)

def segments_to_srt(segments) -> str:
    lines = []
    for i, seg in enumerate(segments, start=1):
        start = format_timestamp(seg.start)
        end = format_timestamp(seg.end)
        text = seg.text.strip()
        lines.append(f"{i}\n{start} --> {end}\n{text}\n")
    return "\n".join(lines)

def write_docx(segments, style="Minimal", include_timestamps=True, title="Transcription") -> bytes:
    doc = Document()
    doc.add_heading(title, 0)
    for seg in segments:
        p = doc.add_paragraph()
        if include_timestamps:
            t = f"[{format_timestamp(seg.start).replace(',', '.')} - {format_timestamp(seg.end).replace(',', '.')}] "
            tr = p.add_run(t)
            tr.bold = True
        # Apply style
        if style == "Minimal":
            r = p.add_run(seg.text.strip())
            r.italic = True
        elif style == "Presentation":
            r = p.add_run(seg.text.strip())
            r.italic = True
        elif style == "Bullet List":
            p.style = "List Bullet"
            r = p.add_run(seg.text.strip())
            r.italic = True
        else:
            p.add_run(seg.text.strip())
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# -------------------------------
# Model loader (with safe fallbacks)
# -------------------------------
@st.cache_resource(show_spinner=False)
def load_model(model_size: str, compute_type: str):
    from faster_whisper import WhisperModel
    tried = set()
    candidates = [compute_type, "int8", "int8_float32", "float32"]
    last_err = None
    for ct in candidates:
        if ct in tried:
            continue
        try:
            return WhisperModel(model_size, compute_type=ct)
        except Exception as e:
            last_err = e
            tried.add(ct)
    raise last_err

# -------------------------------
# Repository helpers
# -------------------------------
def save_record(filename: str, info, segments, style: str, include_timestamps: bool) -> Dict[str, Any]:
    """
    Save transcript outputs to repo/<id>/ and register in index.json.
    """
    ensure_repo()
    rid = new_id()
    item_dir = os.path.join(REPO_DIR, rid)
    os.makedirs(item_dir, exist_ok=True)

    # Build outputs
    plaintext = segments_to_plaintext(segments)
    srt_text = segments_to_srt(segments)
    docx_bytes = write_docx(segments, style=style, include_timestamps=include_timestamps, title="Transcription")

    # Write files
    txt_path = os.path.join(item_dir, "transcript.txt")
    srt_path = os.path.join(item_dir, "subtitles.srt")
    docx_path = os.path.join(item_dir, "transcript.docx")
    meta_path = os.path.join(item_dir, "meta.json")

    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(plaintext)
    with open(srt_path, "w", encoding="utf-8") as f:
        f.write(srt_text)
    with open(docx_path, "wb") as f:
        f.write(docx_bytes)

    meta = {
        "id": rid,
        "filename": filename,
        "created_at": time.time(),
        "language": getattr(info, "language", "unknown"),
        "duration": float(getattr(info, "duration", 0.0)),
        "style": style,
        "include_timestamps": include_timestamps,
        "paths": {
            "txt": txt_path,
            "srt": srt_path,
            "docx": docx_path
        }
    }
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)

    # Update index
    idx = load_index()
    idx["items"].append({
        "id": rid,
        "filename": filename,
        "created_at": meta["created_at"],
        "language": meta["language"],
        "duration": meta["duration"],
        "style": style
    })
    save_index(idx)
    return meta

def read_meta(rid: str) -> Dict[str, Any]:
    meta_path = os.path.join(REPO_DIR, rid, "meta.json")
    with open(meta_path, "r", encoding="utf-8") as f:
        return json.load(f)

def read_text(rid: str) -> str:
    p = os.path.join(REPO_DIR, rid, "transcript.txt")
    with open(p, "r", encoding="utf-8") as f:
        return f.read()

# -------------------------------
# Prompt helpers
# -------------------------------
DEFAULT_PROMPT_TEMPLATE = """You are an expert presentation creator.

Create a presentation outline for a presenter using the following transcript(s).
- Audience: general
- Tone: clear, engaging, concise
- Output: Title, 5-8 key sections with bullet points, and 3 actionable takeaways
- Include data points, quotes, and time markers when useful.

TRANSCRIPTS:
{combined_transcripts}

Guidelines:
- Consolidate duplicate points.
- Keep bullets short and scannable.
- Suggest 3 slide visuals or diagrams.
"""

def build_prompt_for_ids(ids: List[str], template: str) -> str:
    texts = []
    for rid in ids:
        meta = read_meta(rid)
        txt = read_text(rid)
        head = f"\n=== {meta['filename']} (id: {rid}, lang: {meta['language']}, duration: {meta['duration']:.1f}s) ===\n"
        texts.append(head + txt.strip())
    combined = "\n\n".join(texts)
    return template.replace("{combined_transcripts}", combined)

# -------------------------------
# Streamlit UI
# -------------------------------
st.set_page_config(page_title="MP4 → Word Transcriber", page_icon="📝", layout="wide")
st.title("📝 MP4 → Word Transcriber")

with st.sidebar:
    st.header("Settings")
    model_size = st.selectbox(
        "Model",
        ["small.en", "small", "medium.en", "medium", "large-v2"],
        index=0,
        help="*.en = English-only (faster). Non-.en = multilingual."
    )
    compute_type = st.selectbox(
        "Compute type",
        ["float32", "int8", "int8_float32", "float16"],
        index=0,
        help="On CPU use float32/int8; on GPU use float16 if available."
    )
    vad_filter = st.toggle("VAD speech detection", value=True, help="Helps skip long silences/noise.")
    beam_size = st.slider("Beam size", 1, 10, 5)

    st.subheader("DOCX Formatting")
    format_preset = st.selectbox("Style", ["Minimal", "Presentation", "Bullet List"], index=0)
    include_timestamps = st.toggle("Include timecodes in DOCX", value=True)

tabs = st.tabs(["⬆️ Transcribe", "📚 Repository", "✨ Prompt Builder"])

# -------------------------------
# Tab 1: Transcribe
# -------------------------------
with tabs[0]:
    uploaded = st.file_uploader("Upload audio/video file", type=SUPPORTED_TYPES, accept_multiple_files=False)
    if uploaded is not None:
        st.info(f"File: {uploaded.name} • Size: {uploaded.size/1e6:.2f} MB")
        with st.spinner("Preparing audio…"):
            wav_path = tempfile.mktemp(suffix=".wav")
            try:
                extract_audio_to_wav(uploaded.read(), wav_path)
            except RuntimeError as e:
                st.error(str(e))
                st.stop()

        st.success("Audio prepared. Transcribing…")
        from faster_whisper import WhisperModel
        model = load_model(model_size, compute_type)

        segments_out = []
        with st.spinner("Transcribing with Faster-Whisper…"):
            segments, info = model.transcribe(
                wav_path,
                beam_size=beam_size,
                vad_filter=vad_filter,
                language=None,
            )
            for seg in segments:
                segments_out.append(seg)

        try:
            os.unlink(wav_path)
        except OSError:
            pass

        if not segments_out:
            st.error("No speech detected.")
        else:
            left, right = st.columns([2,1])
            with left:
                plaintext = segments_to_plaintext(segments_out)
                st.subheader("Preview transcript")
                st.text_area("Transcript", plaintext, height=320)

            with right:
                srt_text = segments_to_srt(segments_out)
                docx_bytes = write_docx(segments_out, style=format_preset, include_timestamps=include_timestamps)

                st.subheader("Downloads")
                st.download_button("⬇️ Download .txt", data=plaintext, file_name="transcript.txt")
                st.download_button("⬇️ Download .srt", data=srt_text, file_name="subtitles.srt")
                st.download_button("⬇️ Download .docx", data=docx_bytes, file_name="transcript.docx")

                if st.button("💾 Save to repository"):
                    meta = save_record(uploaded.name, info, segments_out, format_preset, include_timestamps)
                    st.success(f"Saved as id: {meta['id']}. Check the Repository tab.")

# -------------------------------
# Tab 2: Repository / Dashboard
# -------------------------------
with tabs[1]:
    ensure_repo()
    idx = load_index()
    rows = idx["items"]
    if not rows:
        st.info("No saved transcripts yet. Save one from the **Transcribe** tab.")
    else:
        df = pd.DataFrame([{
            "Select": False,
            "ID": r["id"],
            "File": r["filename"],
            "Language": r.get("language", "unknown"),
            "Duration (s)": round(r.get("duration", 0.0), 1),
            "Saved": human_dt(r.get("created_at", time.time())),
            "Style": r.get("style", "Minimal"),
        } for r in rows])

        st.subheader("Saved transcripts")
        edited = st.data_editor(
            df,
            num_rows="fixed",
            use_container_width=True,
            column_config={"Select": st.column_config.CheckboxColumn(required=False)},
        )

        st.markdown("### Details & downloads")
        for _, row in edited.iterrows():
            rid = row["ID"]
            with st.expander(f"{row['File']}  —  id: {rid}"):
                meta = read_meta(rid)
                colA, colB, colC = st.columns(3)
                with colA:
                    with open(meta["paths"]["txt"], "r", encoding="utf-8") as f:
                        txt_data = f.read()
                    st.download_button("⬇️ transcript.txt", data=txt_data, file_name=f"{rid}_transcript.txt", use_container_width=True)
                with colB:
                    with open(meta["paths"]["srt"], "r", encoding="utf-8") as f:
                        srt_data = f.read()
                    st.download_button("⬇️ subtitles.srt", data=srt_data, file_name=f"{rid}_subtitles.srt", use_container_width=True)
                with colC:
                    with open(meta["paths"]["docx"], "rb") as f:
                        docx_data = f.read()
                    st.download_button("⬇️ transcript.docx", data=docx_data, file_name=f"{rid}_transcript.docx", use_container_width=True)

                st.caption(f"Saved {human_dt(meta['created_at'])} • lang={meta['language']} • duration={meta['duration']:.1f}s • style={meta['style']}")

                if st.button(f"✨ Create prompt from this ({rid})", key=f"prompt_one_{rid}"):
                    if "prompt_template" not in st.session_state:
                        st.session_state["prompt_template"] = DEFAULT_PROMPT_TEMPLATE
                    prompt = build_prompt_for_ids([rid], st.session_state["prompt_template"])
                    st.session_state["last_prompt"] = prompt
                    st.success("Prompt built. See the **Prompt Builder** tab below.")
                    st.text_area("Prompt preview", prompt, height=200)

        selected_ids = edited[edited["Select"]]["ID"].tolist()
        st.markdown("---")
        st.write("**Batch selection**")
        st.write(f"Selected: {selected_ids or 'None'}")
        if st.button("✨ Build prompt from selected"):
            if not selected_ids:
                st.warning("No items selected.")
            else:
                if "prompt_template" not in st.session_state:
                    st.session_state["prompt_template"] = DEFAULT_PROMPT_TEMPLATE
                prompt = build_prompt_for_ids(selected_ids, st.session_state["prompt_template"])
                st.session_state["last_prompt"] = prompt
                st.success("Prompt built. See the **Prompt Builder** tab.")
                st.text_area("Prompt preview", prompt, height=200)

# -------------------------------
# Tab 3: Prompt Builder
# -------------------------------
with tabs[2]:
    ensure_repo()
    idx = load_index()
    all_items = idx["items"]
    st.subheader("Prompt template")
    template = st.text_area(
        "Edit your template (use {combined_transcripts} where transcripts should be inserted)",
        value=st.session_state.get("prompt_template", DEFAULT_PROMPT_TEMPLATE),
        height=260,
    )
    st.session_state["prompt_template"] = template

    id_to_label = {r["id"]: f"{r['filename']}  (id: {r['id']})" for r in all_items}
    selected = st.multiselect(
        "Pick one or more saved transcripts",
        options=list(id_to_label.keys()),
        format_func=lambda k: id_to_label[k]
    )

    if st.button("✨ Create prompt"):
        if not selected:
            st.warning("Select at least one saved transcript.")
        else:
            prompt = build_prompt_for_ids(selected, template)
            st.session_state["last_prompt"] = prompt
            st.success("Prompt created.")
            st.text_area("Final prompt", prompt, height=300)
            st.download_button("⬇️ Download prompt.txt", data=prompt, file_name="prompt.txt")

st.caption("Note: On Streamlit Cloud, files in the app's filesystem may reset on redeploy or reboot. For permanent storage, we can plug this into S3/Drive/Supabase later.")
